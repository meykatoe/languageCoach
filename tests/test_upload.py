import io

import pytest

import app.routers.upload as upload_module
from app.services.file_extract import UnsupportedFileType, extract_text


def test_extract_text_txt():
    text = extract_text("notes.txt", "Hello world, this is a test.".encode("utf-8"))
    assert text == "Hello world, this is a test."


def test_extract_text_docx_real_file():
    from docx import Document

    doc = Document()
    doc.add_paragraph("First paragraph about grammar.")
    doc.add_paragraph("Second paragraph with a question.")
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_text("essay.docx", buf.getvalue())
    assert "First paragraph about grammar." in text
    assert "Second paragraph with a question." in text


def test_extract_text_unsupported_extension():
    with pytest.raises(UnsupportedFileType):
        extract_text("image.png", b"not really an image")


def test_extract_text_no_extension():
    with pytest.raises(UnsupportedFileType):
        extract_text("noextension", b"content")


def test_upload_rejects_unsupported_type(client):
    res = client.post(
        "/api/upload",
        files={"file": ("photo.png", b"fake bytes", "image/png")},
        data={"count": "2"},
    )
    assert res.status_code == 400


def test_upload_rejects_out_of_range_count(client):
    res = client.post(
        "/api/upload",
        files={"file": ("notes.txt", b"some reference text", "text/plain")},
        data={"count": "20"},
    )
    assert res.status_code == 400


def test_upload_rejects_empty_extracted_text(client):
    res = client.post(
        "/api/upload",
        files={"file": ("empty.txt", b"   \n  ", "text/plain")},
        data={"count": "2"},
    )
    assert res.status_code == 400


def test_upload_creates_questions_from_mocked_ai(client, monkeypatch):
    def fake_generate_from_reference(user_id, reference_text, count):
        assert "vocabulary quiz about animals" in reference_text
        return [
            {
                "id": f"upload-animal-{i}",
                "question": f"Which animal says moo? ({i})",
                "options": ["A. Cow", "B. Cat", "C. Dog", "D. Bird"],
                "answer": "A",
            }
            for i in range(count)
        ]

    monkeypatch.setattr(upload_module, "generate_from_reference", fake_generate_from_reference)

    res = client.post(
        "/api/upload",
        files={"file": ("animals.txt", b"vocabulary quiz about animals", "text/plain")},
        data={"count": "3", "exam": "My Animal Quiz"},
    )
    assert res.status_code == 200
    items = res.json()
    assert len(items) == 3
    assert all(i["source_file"] == "user-upload" for i in items)
    assert all(i["exam"] == "My Animal Quiz" for i in items)
    assert all(i["section"] == "Reading" for i in items)  # multiple-choice shape
    assert all(i["part"] == "animals.txt" for i in items)
    ids = {i["source_id"] for i in items}
    assert len(ids) == 3


def test_upload_writing_prompt_shape_gets_writing_section(client, monkeypatch):
    def fake_generate_from_reference(user_id, reference_text, count):
        return [{"id": "upload-essay-1", "prompt": "Discuss the topic in 250 words."}]

    monkeypatch.setattr(upload_module, "generate_from_reference", fake_generate_from_reference)

    res = client.post(
        "/api/upload",
        files={"file": ("article.txt", b"a long article about climate change", "text/plain")},
        data={"count": "1"},
    )
    assert res.status_code == 200
    items = res.json()
    assert items[0]["section"] == "Writing"
    assert items[0]["exam"] == "Custom"


def test_upload_dedupes_nested_sub_question_ids(client, monkeypatch):
    def fake_generate_from_reference(user_id, reference_text, count):
        return [
            {
                "id": "upload-passage-99",
                "passage": "A new memo about updated parking policies.",
                "questions": [
                    {
                        "id": "toeic-r7-single-01-q1",  # collides with an existing nested id
                        "question": "What is the memo about?",
                        "options": ["A. Parking", "B. Payroll", "C. Holidays", "D. Training"],
                        "answer": "A",
                    }
                ],
            }
        ]

    monkeypatch.setattr(upload_module, "generate_from_reference", fake_generate_from_reference)

    res = client.post(
        "/api/upload",
        files={"file": ("memo.txt", b"a memo about parking", "text/plain")},
        data={"count": "1"},
    )
    assert res.status_code == 200
    nested_id = res.json()[0]["content"]["questions"][0]["id"]
    assert nested_id != "toeic-r7-single-01-q1"

    original = client.post(
        "/api/practice/submit", json={"answers": [{"source_id": "toeic-r7-single-01-q1", "answer": "B"}]}
    )
    assert original.json()["results"][0]["correctAnswer"] == "B"


def test_upload_service_unavailable_returns_503(client, monkeypatch):
    def raise_runtime(*args, **kwargs):
        raise RuntimeError("OpenAI API Key 尚未設定。")

    monkeypatch.setattr(upload_module, "generate_from_reference", raise_runtime)

    res = client.post(
        "/api/upload",
        files={"file": ("notes.txt", b"some text", "text/plain")},
        data={"count": "1"},
    )
    assert res.status_code == 503


def test_upload_created_items_are_fetchable_via_questions_api(client, monkeypatch):
    def fake_generate_from_reference(user_id, reference_text, count):
        return [
            {
                "id": "upload-fetchtest-1",
                "question": "Sample?",
                "options": ["A. Yes", "B. No"],
                "answer": "A",
            }
        ]

    monkeypatch.setattr(upload_module, "generate_from_reference", fake_generate_from_reference)

    res = client.post(
        "/api/upload",
        files={"file": ("fetchtest.txt", b"reference", "text/plain")},
        data={"count": "1", "exam": "FetchTestExam"},
    )
    assert res.status_code == 200

    res2 = client.get("/api/questions", params={"exam": "FetchTestExam", "limit": 10})
    assert res2.status_code == 200
    assert any(q["source_id"] == "upload-fetchtest-1" for q in res2.json())
